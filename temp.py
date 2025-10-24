import shutil
import zipfile
from pathlib import Path
import streamlit as st
import pandas as pd
import os
import datetime as dt
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re


# === Fonctions utiles ===
def date_now():
    mois = ['', 'janvier', 'février', 'mars', 'avril', 'mai', 'juin',
            'juillet', 'août', 'septembre', 'octobre', 'novembre', 'décembre']
    today = dt.date.today()
    return f"{today.day} {mois[today.month]} {today.year}"

def format_nombre(nombre):
    return f"{nombre:,.2f}".replace(',', ' ').replace('.', ',')

def html_to_docx_advanced(html_content, fichier_word, data):
    """
    Convertit le HTML en Word en préservant mieux la mise en forme
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    document = Document()
    
    # Configuration des marges
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # En-tête avec image si disponible
    if os.path.exists('ressources/images/logo.png'):
        try:
            document.add_picture('ressources/images/logo.png', width=Inches(2))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            pass
    
    # Titre principal
    title = document.add_heading(f"Notice d'appel de fonds n°{data['numero_call']}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date et lieu
    p = document.add_paragraph()
    p.add_run(f"{data['date']}\n").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Destinataire
    document.add_paragraph()
    dest = document.add_paragraph()
    dest_run = dest.add_run(f"{data['souscripteur']}\n")
    dest_run.bold = True
    if data['representant']:
        dest.add_run(f"Représenté par : {data['representant']}\n")
    dest.add_run(f"{data['adresse']}\n")
    dest.add_run(f"{data['code_postal']} {data['ville']}\n")
    dest.add_run(f"{data['pays']}")
    
    document.add_paragraph()
    
    # Objet
    objet = document.add_paragraph()
    objet_run = objet.add_run(f"Objet : Appel de fonds n°{data['numero_call']} - {data['nom_fond']}")
    objet_run.bold = True
    objet_run.underline = True
    
    document.add_paragraph()
    
    # Corps du texte
    document.add_paragraph(
        f"Nous avons l'honneur de vous informer qu'un appel de fonds d'un montant total de "
        f"{data['montant_total']} € ({data['pourcentage_call']} % de l'engagement) "
        f"sera effectué le {data['date_call']} sur le fonds {data['nom_fond']}."
    )
    
    if data['texte_fond_couvrir']:
        document.add_paragraph()
        document.add_paragraph("Cet appel de fonds permettra de couvrir :")
        document.add_paragraph(data['texte_fond_couvrir'])
    
    if data['texte_fond_finance']:
        document.add_paragraph()
        document.add_paragraph("Ainsi que de financer :")
        document.add_paragraph(data['texte_fond_finance'])
    
    document.add_paragraph()
    
    # Tableau récapitulatif
    document.add_heading("Votre situation :", level=2)
    
    table = document.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    
    rows_data = [
        ("Montant de l'engagement initial", f"{data['montant_engagement_initial']} €"),
        ("Nombre de parts souscrites", data['nombre_parts_souscrites']),
        ("Catégorie de parts", data['categorie_part']),
        ("Montant à libérer au titre de cet appel", f"{data['montant_a_liberer']} €"),
        ("Total appelé après cet appel", f"{data['total_appele']} €"),
        ("Pourcentage de libération", f"{data['pourcent_liberation']} %"),
        ("Résiduel à appeler", f"{data['residuel']} €")
    ]
    
    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
    
    document.add_paragraph()
    
    # Instructions de virement
    document.add_heading("Modalités de versement :", level=2)
    
    virement = document.add_paragraph()
    virement.add_run("Merci d'effectuer votre virement à l'ordre de :\n\n").bold = True
    
    coord_table = document.add_table(rows=4, cols=2)
    coord_table.style = 'Light List Accent 1'
    
    coord_data = [
        ("Bénéficiaire", data['nom_fond']),
        ("IBAN", "FR76 XXXX XXXX XXXX XXXX XXXX XXX"),
        ("BIC", "XXXXXXXX"),
        ("Référence obligatoire", data['libelle_virement'])
    ]
    
    for i, (label, value) in enumerate(coord_data):
        row = coord_table.rows[i]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value
    
    document.add_paragraph()
    
    # Date limite
    deadline = document.add_paragraph()
    deadline_run = deadline.add_run(f"Date limite de versement : {data['date_call']}")
    deadline_run.bold = True
    deadline_run.font.color.rgb = RGBColor(255, 0, 0)
    deadline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph()
    document.add_paragraph("Nous vous remercions de votre confiance et restons à votre disposition pour tout complément d'information.")
    
    document.add_paragraph()
    signature = document.add_paragraph("Cordialement,\n\nL'équipe de gestion")
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Sauvegarde
    document.save(fichier_word)


# === Interface Streamlit ===
st.title("Générateur de notices d'appel de fonds")

uploaded_file = st.file_uploader("Fichier Excel de données", type=["xlsx"])
texte_fond_couvrir = st.text_area("Texte pour couvrir l'appel")
texte_fond_finance = st.text_area("Texte pour financer le nouvel appel")

numero_call = st.text_input("Numéro de l'appel", value="9")
nom_fond = st.text_input("Nom du fonds", value="FPCI ÉPOPÉE Xplore II")
pays = st.text_input("Pays", value="France")

if st.button("Générer les notices"):
    if uploaded_file:
        # Sauvegarde temporaire
        os.makedirs("ressources", exist_ok=True)
        chemin_fichier = "ressources/Base-data-test-fund-exercice.xlsx"
        with open(chemin_fichier, "wb") as f:
            f.write(uploaded_file.getbuffer())

        try:
            df = pd.read_excel(chemin_fichier, sheet_name='SOUSCRIPTEURS', header=18)
            df_nettoye = df[df['SOUSCRIPTEUR'].notna()]
            df_nettoye = df_nettoye[~df_nettoye['SOUSCRIPTEUR'].str.startswith('TOTAL', na=False)]
            df_nettoye = df_nettoye.reset_index(drop=True)

            df_CALL = pd.read_excel(chemin_fichier, sheet_name='SOUSCRIPTEURS', header=3)
            call = 'CALL #' + numero_call
            montant_total = df[call][df.shape[0]-6]
            date_call = df_CALL.loc[df_CALL['Nominal'] == call, 'Date'].iloc[0]
            pourcentage_call = df_CALL.loc[df_CALL['Nominal'] == call, df_CALL.columns[2]].iloc[0]

            dir = 'ressources'
            env = Environment(loader=FileSystemLoader(dir))
            template = env.get_template('model_notice_img.html')

            os.makedirs('Output', exist_ok=True)
            os.makedirs('Output_HTML', exist_ok=True)
            os.makedirs('Output/PDF', exist_ok=True)
            os.makedirs('Output/Word', exist_ok=True)

            for i in range(df_nettoye.shape[0]):
                total_avant_call = df_nettoye['TOTAL APPELE'][i] - df_nettoye[call][i]
                pourcentage_avant_call = (total_avant_call / df_nettoye['ENGAGEMENT'][i]) * 100

                if pd.isna(df_nettoye["Représentant"][i]):
                    representant = ''
                else:
                    representant = df_nettoye["Représentant"][i]

                data = {
                    'souscripteur': df_nettoye["SOUSCRIPTEUR"][i],
                    'pm_pp': df_nettoye["TYPE"][i],
                    'representant': representant,
                    'adresse': df_nettoye["ADRESSE"][i],
                    'code_postal': round(df_nettoye["CP"][i]),
                    'ville': df_nettoye["VILLE"][i],
                    'pays': pays,
                    'date': date_now(),
                    'numero_call': numero_call,
                    'date_call': date_call.strftime('%d/%m/%Y'),
                    'nom_fond': nom_fond,
                    'montant_total': format_nombre(montant_total),
                    'pourcentage_call': f"{pourcentage_call * 100:.2f}",
                    'montant_a_liberer': format_nombre(df_nettoye[call][i]),
                    'pourcentage_avant_call': format_nombre(pourcentage_avant_call),
                    'texte_fond_couvrir': texte_fond_couvrir,
                    'texte_fond_finance': texte_fond_finance,
                    'montant_engagement_initial': format_nombre(df_nettoye["ENGAGEMENT"][i]),
                    'nombre_parts_souscrites': format_nombre(df_nettoye["NBR PARTS"][i]),
                    'categorie_part': df_nettoye["PART"][i],
                    'total_appele': format_nombre(df_nettoye["TOTAL APPELE"][i]),
                    'pourcent_liberation': f"{df_nettoye['%LIBERATION'][i] * 100:.2f}",
                    'residuel': format_nombre(df_nettoye["RESIDUEL"][i]),
                    'libelle_virement': 'CR ' + df_nettoye["SOUSCRIPTEUR"][i] + ' ADF ' + numero_call
                }

                # Génération HTML
                html_content = template.render(data)
                dir_nom_fichier = f'Output_HTML/{df_nettoye["SOUSCRIPTEUR"][i]}_{df_nettoye["PART"][i]}.html'
                with open(dir_nom_fichier, 'w', encoding='utf-8') as file:
                    file.write(html_content)

                # Génération PDF
                fichier_html = f'Output_HTML/{df_nettoye["SOUSCRIPTEUR"][i]}_{df_nettoye["PART"][i]}.html'
                fichier_pdf = f'Output/PDF/{df_nettoye["SOUSCRIPTEUR"][i]}_{df_nettoye["PART"][i]}.pdf'
                base_url = Path('ressources/images').resolve()
                HTML(filename=fichier_html, base_url=base_url.as_uri()).write_pdf(fichier_pdf)

                # Génération Word AMÉLIORÉE
                fichier_word = f'Output/Word/{df_nettoye["SOUSCRIPTEUR"][i]}_{df_nettoye["PART"][i]}.docx'
                html_to_docx_advanced(html_content, fichier_word, data)

            # Zip tous les fichiers
            shutil.make_archive("notices", "zip", "Output")
            with open("notices.zip", "rb") as f:
                st.download_button("Télécharger les notices générées", f, "notices.zip")
            
            st.success(f"✅ {df_nettoye.shape[0]} notices générées avec succès !")

        except Exception as e:
            st.error(f"Une erreur est survenue : {e}")
            import traceback
            st.code(traceback.format_exc())
    else:
        st.warning("Merci de déposer un fichier Excel avant de lancer la génération.")